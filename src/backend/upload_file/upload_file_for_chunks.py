import logging
import json
import os
from minio import Minio
from minio.error import S3Error
from fastapi import HTTPException
import PyPDF2
import pdfplumber
from docx import Document

# Logger setup
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

class DocumentReader:
    def __init__(self):

        try:
            self.client = Minio(
                os.getenv('MINIO_HOST', 'minio:9000'),
                access_key=os.getenv('MINIO_ROOT_USER', 'minioadmin'),
                secret_key=os.getenv('MINIO_ROOT_PASSWORD', 'minioadmin'),
                secure=False
            )
            logger.info("MinIO client initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing MinIO client: {e}")
            raise HTTPException(status_code=500, detail="Error initializing MinIO client")

    def upload_file(self, bucket_name, filename):

        try:
            self.client.put_object(
                bucket_name=bucket_name,
                object_name=filename,
            )
            logger.info(f"File '{filename}' uploaded successfully to bucket '{bucket_name}'")
            return bucket_name, filename
        except S3Error as e:
            logger.error(f"Error uploading file '{filename}' to bucket '{bucket_name}': {e}")
            raise HTTPException(status_code=500, detail="Error uploading file")

    def get_customer_guid_and_filename(self):
        customer_guid = input("Enter the customer GUID: ")
        filename = input("Enter the filename (with extension): ")
        return self.upload_file(customer_guid, filename)

    def get_file_from_minio(self, filename):
        try:
            file_data = self.client.get_object(filename)
            return file_data
        except S3Error as e:
            logger.error(f"Error retrieving file '{filename}' from MinIO: {e}")
            raise HTTPException(status_code=500, detail="Error retrieving file")

    def check_file_type(self, filename):
        approaches = []
        file_type = None

        # Approach 1: Check extension
        if filename.endswith(".pdf"):
            approaches.append("Approach 1: File is a PDF based on extension")
            file_type = "pdf"
        elif filename.endswith(".docx"):
            approaches.append("Approach 2: File is a DOCX based on extension")
            file_type = "docx"
        elif filename.endswith(".txt"):
            approaches.append("Approach 3: File is a TXT based on extension")
            file_type = "txt"
        else:
            approaches.append("No match: Unsupported file extension")

        # Approach 2: Check content by file signature (simplified)
        if filename.endswith(".pdf"):
            try:
                with open(filename, "rb") as f:
                    header = f.read(4)
                    if header == b"%PDF":
                        approaches.append("Approach 2: PDF signature matched")
            except Exception:
                approaches.append("Approach 2: PDF signature not matched")

        elif filename.endswith(".docx"):
            try:
                with open(filename, "rb") as f:
                    header = f.read(4)
                    if header == b"PK\x03\x04":
                        approaches.append("Approach 2: DOCX signature matched")
            except Exception:
                approaches.append("Approach 2: DOCX signature not matched")

        elif filename.endswith(".txt"):
            try:
                with open(filename, "r") as f:
                    content = f.read(100)
                    if content.isprintable():
                        approaches.append("Approach 2: TXT content matched")
            except Exception:
                approaches.append("Approach 2: TXT content not matched")

        # Approach 3: Check using libraries
        if filename.endswith(".pdf"):
            try:
                with open(filename, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    if len(reader.pages) > 0:
                        approaches.append("Approach 3: PDF content extracted")
            except Exception:
                approaches.append("Approach 3: PDF content extraction failed")

        elif filename.endswith(".docx"):
            try:
                doc = Document(filename)
                if len(doc.paragraphs) > 0:
                    approaches.append("Approach 3: DOCX content extracted")
            except Exception:
                approaches.append("Approach 3: DOCX content extraction failed")

        elif filename.endswith(".txt"):
            try:
                with open(filename, "r") as f:
                    content = f.read(100)
                    if content:
                        approaches.append("Approach 3: TXT content validated")
            except Exception:
                approaches.append("Approach 3: TXT content validation failed")

        #Select valid approaches
        valid_approaches = [approach for approach in approaches if "matched" in approach or "extracted" in approach]
        if len(valid_approaches) >= 2:
            return file_type
        else:
            raise HTTPException(status_code=400, detail="File type mismatch, extraction failed.")

    def extract_pdf_content(self, file_data):

        try:
            with pdfplumber.open(file_data) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text()
            return text
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {e}")
            raise HTTPException(status_code=400, detail="Failed to extract text from PDF")

    def extract_docx_content(self, file_data):

        try:
            doc = Document(file_data)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            raise HTTPException(status_code=400, detail="Failed to extract text from DOCX")

    def extract_txt_content(self, file_data):

        try:
            text = file_data.decode('utf-8')
            return text
        except Exception as e:
            logging.error(f"Error extracting text from TXT: {e}")
            raise HTTPException(status_code=400, detail="Failed to extract text from TXT")

    def extract_and_read_file(self, file_data, file_format):

        if file_format == 'pdf':
            return self.extract_pdf_content(file_data)
        elif file_format == 'docx':
            return self.extract_docx_content(file_data)
        elif file_format == 'txt':
            return self.extract_txt_content(file_data)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")

    def process_document(self):

        #Get customer GUID and filename from user input
        customer_guid, filename = self.get_customer_guid_and_filename()

        #Retrieve the document from MinIO
        file_data = self.get_file_from_minio(filename)

        #Check file type and validation approaches
        file_type = self.check_file_type(filename)

        #Extract content based on the file format
        extracted_content = self.extract_and_read_file(file_data, file_type)

        #Convert to JSON format
        file_json = {
            "content": extracted_content
        }
        return json.dumps(file_json)

# The program starts here
if __name__ == "__main__":
    doc_reader = DocumentReader()

    #Get customer GUID and filename manually from user input
    json_output = doc_reader.process_document()

    #Output the extracted content as JSON
    print(json_output)
